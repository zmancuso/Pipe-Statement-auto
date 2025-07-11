from flask import Flask, render_template, request, send_file
from docx import Document
from io import BytesIO

app = Flask(__name__)

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        data = request.form
        doc = Document()

        # Title
        doc.add_heading("Customer Account Statement", 0)

        # Account Summary
        doc.add_heading("Account Summary", level=1)
        doc.add_paragraph(f"Collection Start Date: {data['collectionStart']}")
        doc.add_paragraph(f"Total Amount Advanced: {data['totalAdvanced']}")
        doc.add_paragraph(f"Advance Count: {data['advanceCount']}")
        doc.add_paragraph(f"Total Fee: {data['totalFee']}")
        doc.add_paragraph(f"Total Obligation: {data['totalObligation']}")
        doc.add_paragraph(f"Total Outstanding Balance: {data['outstandingBalance']}")

        # Payment Performance
        doc.add_heading("Payment Performance", level=1)
        doc.add_paragraph(f"Successful Payments: {data['successfulPayments']}")
        doc.add_paragraph(f"Failed Payments: {data['failedPayments']}")

        # Failed Payment Details
        doc.add_heading("Failed Payment Details", level=1)
        failed_table = doc.add_table(rows=1, cols=4)
        failed_table.style = "Table Grid"
        failed_hdr = failed_table.rows[0].cells
        failed_hdr[0].text = "Date"
        failed_hdr[1].text = "Amount"
        failed_hdr[2].text = "Status"
        failed_hdr[3].text = "Reason"

        for line in data['failedDetails'].strip().split("\n"):
            parts = [p.strip() for p in line.split(",")]
            if len(parts) == 4:
                row_cells = failed_table.add_row().cells
                for i in range(4):
                    row_cells[i].text = parts[i]

        # Recent Payment Activity
        doc.add_heading("Recent Payment Activity", level=1)
        rev_table = doc.add_table(rows=1, cols=10)
        rev_table.style = "Table Grid"
        headers = [
            "Revenue Date", "Revenue", "Collected", "Method",
            "Collection Date", "Source", "Increase", "Status",
            "External Link", "Attempts"
        ]
        for i, h in enumerate(headers):
            rev_table.rows[0].cells[i].text = h

        lines = data["revenueHistory"].strip().split("\n")
        for i in range(0, len(lines), 10):
            row = lines[i:i + 10]
            while len(row) < 10:
                row.append("")  # pad missing cells
            row_cells = rev_table.add_row().cells
            for j in range(10):
                row_cells[j].text = row[j].strip()

        # Next Steps & Notes
        doc.add_heading("Next Steps & Notes", level=1)
        doc.add_paragraph("Failed Payment Balance remains to be collected. Please let us know if and when we can resubmit failed payments as soon as possible to ensure the account is in good standing.")
        doc.add_paragraph(f"Outstanding Balance: {data['outstandingBalance']} remains to be collected.")
        doc.add_paragraph("If you have any questions or disputes regarding your balance, please contact the Pipe Servicing & Collections Team at collections@pipe.com or +1 845-704-4476.")

        # Stream DOCX to browser
        f = BytesIO()
        doc.save(f)
        f.seek(0)
        return send_file(f, as_attachment=True, download_name="Customer_Statement.docx")

    return render_template("form.html")
